from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def render_docx(html_content: str, output_path: str = "resume.docx") -> str:
    """
    Convert HTML resume to DOCX format
    
    Args:
        html_content: HTML string of the resume
        output_path: Path to save the DOCX file
        
    Returns:
        Path to the generated DOCX file
    """
    # Create a new document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Resume"
    doc.core_properties.author = "AI Resume Builder"
    
    # Add styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Create heading style
    heading_style = doc.styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Create subheading style
    subheading_style = doc.styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.size = Pt(14)
    subheading_style.font.bold = True
    subheading_style.font.color.rgb = RGBColor(44, 62, 80)
    
    # Create section style
    section_style = doc.styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.size = Pt(12)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(52, 73, 94)
    
    # Extract text from HTML (simplified parsing)
    # In a real implementation, you would use BeautifulSoup or similar
    text_content = extract_text_from_html(html_content)
    
    # Parse and add content to document
    sections = text_content.split('\n\n')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
            
        # Check if it's a heading
        if section.upper() in ['PROFESSIONAL SUMMARY', 'EXPERIENCE', 'EDUCATION', 
                              'SKILLS', 'PROJECTS', 'ACHIEVEMENTS']:
            p = doc.add_paragraph(section.upper())
            p.style = 'Heading2'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # Check if it's a name/title
        elif 'Resume' in section or len(section.split()) <= 3:
            p = doc.add_paragraph(section)
            p.style = 'Heading1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        else:
            # Check for bullet points
            if '•' in section or '-' in section:
                lines = section.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith(('•', '-', '*')):
                        line = line[1:].strip()
                        p = doc.add_paragraph(line, style='List Bullet')
                    elif line:
                        p = doc.add_paragraph(line)
            else:
                # Regular paragraph
                p = doc.add_paragraph(section)
    
    # Save document
    doc.save(output_path)
    return output_path

def extract_text_from_html(html: str) -> str:
    """Extract clean text from HTML"""
    # Remove script and style elements
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
    
    # Replace common HTML elements
    html = re.sub(r'<br\s*/?>', '\n', html)
    html = re.sub(r'</p>', '\n\n', html)
    html = re.sub(r'</div>', '\n', html)
    html = re.sub(r'</h[1-6]>', '\n\n', html)
    html = re.sub(r'<li>', '• ', html)
    html = re.sub(r'</li>', '\n', html)
    
    # Remove all HTML tags
    text = re.sub(r'<[^>]+>', '', html)
    
    # Clean up whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Decode HTML entities (simplified)
    replacements = {
        '&nbsp;': ' ',
        '&amp;': '&',
        '&lt;': '<',
        '&gt;': '>',
        '&quot;': '"',
        '&#39;': "'"
    }
    
    for entity, replacement in replacements.items():
        text = text.replace(entity, replacement)
    
    return text.strip()
